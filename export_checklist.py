#!/usr/bin/env python3
"""匯出觀察清單的檢核表為 Word 檔"""
import sqlite3, json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

DB_PATH = os.path.join(os.path.dirname(__file__), 'stocks.db')

# 檢核項目定義（與 app.py CHECKLIST_ITEMS 一致）
CHECKLIST_ITEMS = [
    {'key': 'roic_avg5',      'category': 'profit', 'label': 'ROIC 近5年平均 ≥15%'},
    {'key': 'roic_latest',    'category': 'profit', 'label': 'ROIC 最近一年 ≥15%'},
    {'key': 'roic_min5',      'category': 'profit', 'label': 'ROIC 近5年最低值 ≥10%'},
    {'key': 'opm_avg5',       'category': 'profit', 'label': '營益率近5年平均 ≥10%'},
    {'key': 'opm_min5',       'category': 'profit', 'label': '營益率近5年最低值 ≥5%'},
    {'key': 'gm_median',      'category': 'profit', 'label': '毛利率 ≥ 近5年中位數'},
    {'key': 'gm_q_median',    'category': 'profit', 'label': '最近一季毛利率 ≥ 近4季中位數'},
    {'key': 'debt_ratio_ok',  'category': 'safety', 'label': '負債比 ≤ 50%'},
    {'key': 'fin_debt_ok',    'category': 'safety', 'label': '金融負債比 < 30%'},
    {'key': 'icr_ok',         'category': 'safety', 'label': '利息保障倍數 > 5'},
    {'key': 'fcf_freq',       'category': 'safety', 'label': 'FCF近5年至少3年為正'},
    {'key': 'fcf_no_consec',  'category': 'safety', 'label': 'FCF近2年不得連續為負'},
    {'key': 'fcf_sum_pos',    'category': 'safety', 'label': 'FCF近5年加總為正'},
    {'key': 'inv_level',      'category': 'safety', 'label': '存貨水準 ≤ 近5年平均×1.2'},
    {'key': 'inv_trend',      'category': 'safety', 'label': '存貨方向：最近一季 ≤ 近4季中位數×1.15'},
    {'key': 'ar_level',       'category': 'safety', 'label': '應收水準 ≤ 近5年平均×1.2'},
    {'key': 'ar_trend',       'category': 'safety', 'label': '應收方向：最近一季 ≤ 近4季中位數×1.15'},
    {'key': 'grade_a_ok',      'category': 'value', 'label': '預估(沈董)等級為A級以上'},
    {'key': 'eps_vs_median5',  'category': 'value', 'label': '預估(沈董)EPS >= 近5年EPS中位數'},
    {'key': 'core_ratio',      'category': 'value', 'label': '累計營業利益 / 累計稅前淨利 > 70%'},
    {'key': 'price_val_ok',    'category': 'value', 'label': '現價 <= A級評價；<= AA更佳'},
    {'key': 'val_ddm_return', 'category': 'value', 'label': '股利折現現價潛在年報酬 >= 10%'},
    {'key': 'dcf_safe_ok',    'category': 'value', 'label': '現價 <= DCF安全邊際價'},
    {'key': 'rev_cagr5_ok',   'category': 'growth_eval', 'label': '近5年營收CAGR > 5%'},
    {'key': 'eps_cagr5_ok',   'category': 'growth_eval', 'label': '近5年保守成長率(EPS) > 5%'},
    {'key': 'rev_accel',      'category': 'growth_eval', 'label': '近3年營收CAGR > 近5年營收CAGR'},
    {'key': 'eps_accel',      'category': 'growth_eval', 'label': '近3年EPS CAGR > 近5年EPS CAGR'},
    {'key': 'cum_rev_pos',    'category': 'growth_eval', 'label': '累積營收年增率 >= 0%'},
    {'key': 'rev_3m_pos',     'category': 'growth_eval', 'label': '短期3M營收年增率 >= 0%'},
    {'key': 'rev_12m_pos',    'category': 'growth_eval', 'label': '長期12M營收年增率 >= 0%'},
    {'key': 'rev_3m_gt_12m',  'category': 'growth_eval', 'label': '短期3M >= 長期12M'},
    {'key': 'ge_neff_ratio',   'category': 'value', 'label': 'Neff 比率 >= 1.0'},
    {'key': 'ge_growth_green', 'category': 'growth_eval', 'label': '趨勢燈號為多頭（3M/12M+EPS綜合）'},
]

CAT_LABELS = {
    'profit': '獲利性檢核',
    'safety': '安全性檢核',
    'value': '價值評估檢核',
    'growth_eval': '成長性評估檢核',
}
CAT_ORDER = ['profit', 'safety', 'value', 'growth_eval']

def set_cell(cell, text, bold=False, size=8, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')

def main():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row

    # 取觀察清單
    watch_codes = [r['code'] for r in conn.execute(
        "SELECT code FROM user_lists WHERE list_type='watch' ORDER BY code"
    )]
    if not watch_codes:
        print("觀察清單為空")
        return

    # 取股票名稱與收盤價
    placeholders = ','.join('?' * len(watch_codes))
    stocks = {}
    for r in conn.execute(
        f"SELECT code, name, close, shen_eps, shen_grade, val_aa, val_a1, val_a2, val_a "
        f"FROM stocks WHERE code IN ({placeholders})", watch_codes
    ):
        stocks[r['code']] = dict(r)

    # 取檢核表
    checklists = {}
    for r in conn.execute(
        f"SELECT * FROM stock_checklist WHERE code IN ({placeholders})", watch_codes
    ):
        checklists[r['code']] = dict(r)

    conn.close()

    # 建立 Word
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(9)

    # 標題
    title = doc.add_heading('觀察清單 - 檢核表', level=1)
    title.runs[0].font.size = Pt(16)
    p = doc.add_paragraph(f'匯出日期：{datetime.now().strftime("%Y/%m/%d %H:%M")}　共 {len(watch_codes)} 支')
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    for code in watch_codes:
        st = stocks.get(code, {})
        cl = checklists.get(code, {})
        name = st.get('name', '')
        close = st.get('close')
        shen_eps = st.get('shen_eps')
        shen_grade = st.get('shen_grade', '')

        # 股票標題
        doc.add_paragraph()  # 間距
        h = doc.add_heading(level=2)
        run = h.add_run(f'{code} {name}')
        run.font.size = Pt(13)

        # 摘要
        parts = []
        if close is not None: parts.append(f'收盤價 {close:.2f}')
        if shen_eps is not None: parts.append(f'沈董EPS {shen_eps:.2f}')
        if shen_grade: parts.append(f'等級 {shen_grade}')
        val_aa = st.get('val_aa')
        val_a = st.get('val_a')
        if val_aa is not None: parts.append(f'AA={val_aa:.2f}')
        if val_a is not None: parts.append(f'A={val_a:.2f}')

        # 通過統計
        profit_count = cl.get('profit_count', 0) or 0
        safety_count = cl.get('safety_count', 0) or 0
        value_count = cl.get('value_count', 0) or 0
        growth_count = cl.get('growth_eval_count', 0) or 0

        if parts:
            p = doc.add_paragraph('　'.join(parts))
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(71, 85, 105)

        # 按類別分表
        detail = {}
        if cl.get('detail'):
            try:
                detail = json.loads(cl['detail'])
            except Exception:
                pass

        for cat in CAT_ORDER:
            items = [it for it in CHECKLIST_ITEMS if it['category'] == cat]
            if not items:
                continue

            cat_label = CAT_LABELS[cat]
            counts = {
                'profit': profit_count, 'safety': safety_count,
                'value': value_count, 'growth_eval': growth_count
            }
            cat_count = counts.get(cat, 0)

            # 表格：# | 項目 | 是 | 否
            table = doc.add_table(rows=1 + len(items), cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            # 表頭
            hdr = table.rows[0].cells
            set_cell(hdr[0], '#', bold=True, size=7)
            set_cell(hdr[1], f'{cat_label}（{cat_count}/{len(items)}）', bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell(hdr[2], 'V', bold=True, size=7, color=RGBColor(22, 163, 74))
            set_cell(hdr[3], 'X', bold=True, size=7, color=RGBColor(220, 38, 38))
            for c in hdr:
                shade_cell(c, 'E2E8F0')

            # 欄寬
            table.columns[0].width = Cm(0.8)
            table.columns[1].width = Cm(12)
            table.columns[2].width = Cm(0.8)
            table.columns[3].width = Cm(0.8)

            for i, it in enumerate(items):
                row = table.rows[i + 1].cells
                chk_val = cl.get(f'chk_{it["key"]}')
                passed = chk_val == 1
                d = detail.get(it['key'], '')

                label_text = f'{it["label"]}'
                if d:
                    label_text += f'\n{d}'

                set_cell(row[0], str(i + 1), size=7)
                set_cell(row[1], label_text, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
                if passed:
                    set_cell(row[2], 'V', size=7, color=RGBColor(22, 163, 74))
                    set_cell(row[3], '', size=7)
                else:
                    set_cell(row[2], '', size=7)
                    set_cell(row[3], 'X', size=7, color=RGBColor(220, 38, 38))

    # 儲存
    output = '/Users/roger/Documents/AI機器人/觀察清單_檢核表.docx'
    doc.save(output)
    print(f'已匯出：{output}（{len(watch_codes)} 支）')

if __name__ == '__main__':
    main()
